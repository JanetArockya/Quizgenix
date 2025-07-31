import os
import tempfile
from fpdf import FPDF
import pandas as pd
from docx import Document
from datetime import datetime

class FileService:
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def generate_pdf(self, quiz_data):
        """Generate PDF file for quiz"""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('Arial', 'B', 16)
            
            # Title
            pdf.cell(0, 10, f"Quiz: {quiz_data['title']}", ln=True, align='C')
            pdf.ln(5)
            
            # Quiz info
            pdf.set_font('Arial', '', 12)
            pdf.cell(0, 8, f"Subject: {quiz_data['subject']}", ln=True)
            pdf.cell(0, 8, f"Topic: {quiz_data.get('topic', 'N/A')}", ln=True)
            pdf.cell(0, 8, f"Difficulty: {quiz_data['difficulty'].capitalize()}", ln=True)
            pdf.cell(0, 8, f"Total Questions: {len(quiz_data['questions'])}", ln=True)
            pdf.ln(10)
            
            # Questions
            for i, question in enumerate(quiz_data['questions'], 1):
                pdf.set_font('Arial', 'B', 12)
                pdf.cell(0, 8, f"Question {i}:", ln=True)
                
                pdf.set_font('Arial', '', 11)
                # Handle long questions with multi_cell
                pdf.multi_cell(0, 6, question['question'])
                pdf.ln(2)
                
                # Options
                for j, option in enumerate(question['options']):
                    option_letter = chr(65 + j)  # A, B, C, D
                    is_correct = option == question['correct_answer']
                    marker = " ✓" if is_correct else ""
                    pdf.cell(0, 5, f"   {option_letter}) {option}{marker}", ln=True)
                
                pdf.ln(5)
            
            # Save file
            filename = f"quiz_{quiz_data['title'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(self.temp_dir, filename)
            pdf.output(file_path)
            
            return file_path
            
        except Exception as e:
            print(f"Error generating PDF: {e}")
            raise
    
    def generate_word(self, quiz_data):
        """Generate Word document for quiz"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(f"Quiz: {quiz_data['title']}", 0)
            title.alignment = 1  # Center alignment
            
            # Quiz info
            doc.add_paragraph(f"Subject: {quiz_data['subject']}")
            doc.add_paragraph(f"Topic: {quiz_data.get('topic', 'N/A')}")
            doc.add_paragraph(f"Difficulty: {quiz_data['difficulty'].capitalize()}")
            doc.add_paragraph(f"Total Questions: {len(quiz_data['questions'])}")
            doc.add_paragraph()
            
            # Questions
            for i, question in enumerate(quiz_data['questions'], 1):
                # Question heading
                doc.add_heading(f"Question {i}:", level=2)
                
                # Question text
                doc.add_paragraph(question['question'])
                
                # Options
                for j, option in enumerate(question['options']):
                    option_letter = chr(65 + j)  # A, B, C, D
                    is_correct = option == question['correct_answer']
                    marker = " ✓ (Correct Answer)" if is_correct else ""
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{option_letter}) ").bold = True
                    p.add_run(f"{option}{marker}")
                
                doc.add_paragraph()
            
            # Save file
            filename = f"quiz_{quiz_data['title'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = os.path.join(self.temp_dir, filename)
            doc.save(file_path)
            
            return file_path
            
        except Exception as e:
            print(f"Error generating Word document: {e}")
            raise
    
    def generate_excel(self, quiz_data):
        """Generate Excel file for quiz"""
        try:
            # Prepare data for Excel
            data = []
            
            for i, question in enumerate(quiz_data['questions'], 1):
                row = {
                    'Question_Number': i,
                    'Question': question['question'],
                    'Option_A': question['options'][0] if len(question['options']) > 0 else '',
                    'Option_B': question['options'][1] if len(question['options']) > 1 else '',
                    'Option_C': question['options'][2] if len(question['options']) > 2 else '',
                    'Option_D': question['options'][3] if len(question['options']) > 3 else '',
                    'Correct_Answer': question['correct_answer'],
                    'Subject': quiz_data['subject'],
                    'Topic': quiz_data.get('topic', ''),
                    'Difficulty': quiz_data['difficulty']
                }
                data.append(row)
            
            # Create DataFrame
            df = pd.DataFrame(data)
            
            # Save file
            filename = f"quiz_{quiz_data['title'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            file_path = os.path.join(self.temp_dir, filename)
            
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                # Quiz info sheet
                info_df = pd.DataFrame([
                    ['Quiz Title', quiz_data['title']],
                    ['Subject', quiz_data['subject']],
                    ['Topic', quiz_data.get('topic', 'N/A')],
                    ['Difficulty', quiz_data['difficulty'].capitalize()],
                    ['Total Questions', len(quiz_data['questions'])],
                    ['Generated On', datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
                ], columns=['Property', 'Value'])
                
                info_df.to_excel(writer, sheet_name='Quiz_Info', index=False)
                
                # Questions sheet
                df.to_excel(writer, sheet_name='Questions', index=False)
            
            return file_path
            
        except Exception as e:
            print(f"Error generating Excel file: {e}")
            raise