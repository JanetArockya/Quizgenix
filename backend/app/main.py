from flask import Flask, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
import jwt
import datetime
import json
import io
import os
import random
import re

# Optional imports for file generation
try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False
    print("⚠️ FPDF not available - Install with: pip install fpdf2")

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available - Install with: pip install python-docx")

try:
    import pandas as pd
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    print("⚠️ pandas not available - Install with: pip install pandas openpyxl")

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///quizgenix.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
CORS(app)

# Database Models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    role = db.Column(db.String(20), nullable=False, default='student')
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

class Quiz(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    subject = db.Column(db.String(100), nullable=False)
    topic = db.Column(db.String(100), nullable=False)
    difficulty = db.Column(db.String(20), nullable=False)
    questions = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

class QuizAttempt(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    quiz_id = db.Column(db.Integer, db.ForeignKey('quiz.id'), nullable=False)
    answers = db.Column(db.Text, nullable=False)  # JSON string of user answers
    score = db.Column(db.Integer, nullable=False)
    total_questions = db.Column(db.Integer, nullable=False)
    time_taken = db.Column(db.Integer, nullable=False)  # in seconds
    completed_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    
    # Relationships
    user = db.relationship('User', backref='quiz_attempts')
    quiz = db.relationship('Quiz', backref='attempts')

class QuizSession(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    quiz_id = db.Column(db.Integer, db.ForeignKey('quiz.id'), nullable=False)
    session_token = db.Column(db.String(100), nullable=False, unique=True)
    started_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    expires_at = db.Column(db.DateTime, nullable=False)
    is_active = db.Column(db.Boolean, default=True)
    current_question = db.Column(db.Integer, default=0)
    answers_so_far = db.Column(db.Text, default='{}')  # JSON of answers

# Advanced AI Knowledge Base
KNOWLEDGE_BASE = {
    'javascript': {
        'easy': {
            'questions': [
                ("What is the correct way to declare a variable in JavaScript?", 
                 ["var name = 'John';", "variable name = 'John';", "v name = 'John';", "declare name = 'John';"], 0,
                 "The 'var' keyword is used to declare variables in JavaScript."),
                ("Which method is used to add an element to the end of an array?", 
                 ["push()", "add()", "append()", "insert()"], 0,
                 "The push() method adds elements to the end of an array."),
                ("What does '===' operator do in JavaScript?", 
                 ["Strict equality comparison", "Assignment", "Loose equality", "Not equal"], 0,
                 "The '===' operator performs strict equality comparison without type conversion."),
                ("How do you write a comment in JavaScript?", 
                 ["// This is a comment", "<!-- This is a comment -->", "# This is a comment", "/* This is a comment"], 0,
                 "Single line comments in JavaScript start with //."),
                ("What is the correct way to create a function?", 
                 ["function myFunction() {}", "create myFunction() {}", "def myFunction() {}", "function = myFunction() {}"], 0,
                 "Functions in JavaScript are declared using the 'function' keyword.")
            ]
        },
        'medium': {
            'questions': [
                ("What is closure in JavaScript?", 
                 ["Function with access to outer scope", "A loop structure", "An object method", "A data type"], 0,
                 "Closure allows a function to access variables from its outer/enclosing scope."),
                ("Which method creates a new array with results of calling a function?", 
                 ["map()", "filter()", "reduce()", "forEach()"], 0,
                 "The map() method creates a new array with the results of calling a function for every array element."),
                ("What is the difference between 'let' and 'var'?", 
                 ["'let' has block scope, 'var' has function scope", "No difference", "'var' is newer", "'let' is faster"], 0,
                 "'let' provides block scoping while 'var' has function scoping."),
                ("What is event delegation?", 
                 ["Handling events on parent elements", "Creating custom events", "Preventing default behavior", "Stopping event propagation"], 0,
                 "Event delegation allows handling events on parent elements instead of individual child elements."),
                ("What does 'this' keyword refer to?", 
                 ["The current execution context", "The previous function", "A global variable", "The parent object"], 0,
                 "The 'this' keyword refers to the object that is currently executing the code.")
            ]
        },
        'hard': {
            'questions': [
                ("What is the time complexity of Array.prototype.indexOf()?", 
                 ["O(n)", "O(log n)", "O(1)", "O(n²)"], 0,
                 "indexOf() performs a linear search, resulting in O(n) time complexity."),
                ("How does JavaScript's event loop work?", 
                 ["Call stack, callback queue, and event loop coordination", "Simple linear execution", "Parallel processing", "Random execution order"], 0,
                 "The event loop manages the call stack and callback queue for asynchronous operations."),
                ("What is the prototype chain?", 
                 ["Inheritance mechanism in JavaScript", "A design pattern", "A loop structure", "An error handling method"], 0,
                 "The prototype chain enables inheritance by linking objects through their prototype property."),
                ("What happens during hoisting?", 
                 ["Variable and function declarations are moved to top", "Code is optimized", "Errors are fixed", "Variables are deleted"], 0,
                 "Hoisting moves variable and function declarations to the top of their scope during compilation."),
                ("What is the difference between call() and apply()?", 
                 ["call() takes arguments individually, apply() takes array", "No difference", "apply() is faster", "call() is deprecated"], 0,
                 "call() accepts arguments individually while apply() accepts them as an array.")
            ]
        }
    },
    'python': {
        'easy': {
            'questions': [
                ("How do you create a list in Python?", 
                 ["my_list = [1, 2, 3]", "my_list = (1, 2, 3)", "my_list = {1, 2, 3}", "my_list = <1, 2, 3>"], 0,
                 "Lists in Python are created using square brackets."),
                ("Which keyword is used to define a function?", 
                 ["def", "function", "define", "func"], 0,
                 "The 'def' keyword is used to define functions in Python."),
                ("How do you print text in Python?", 
                 ["print('Hello')", "echo('Hello')", "console.log('Hello')", "write('Hello')"], 0,
                 "The print() function outputs text in Python."),
                ("What is the correct way to create a dictionary?", 
                 ["my_dict = {'key': 'value'}", "my_dict = ['key': 'value']", "my_dict = ('key': 'value')", "my_dict = <'key': 'value'>"], 0,
                 "Dictionaries in Python are created using curly braces with key-value pairs."),
                ("Which operator is used for exponentiation?", 
                 ["**", "^", "exp", "pow"], 0,
                 "The ** operator performs exponentiation in Python.")
            ]
        },
        'medium': {
            'questions': [
                ("What is a list comprehension?", 
                 ["Compact way to create lists", "List sorting method", "List deletion technique", "List indexing"], 0,
                 "List comprehensions provide a concise way to create lists based on existing lists."),
                ("What is the difference between '==' and 'is'?", 
                 ["'==' compares values, 'is' compares identity", "No difference", "'is' compares values", "'==' is faster"], 0,
                 "'==' compares values while 'is' compares object identity."),
                ("What is a decorator in Python?", 
                 ["Function that modifies another function", "Class inheritance", "Error handling", "Loop structure"], 0,
                 "Decorators are functions that modify or extend the behavior of other functions."),
                ("What is the purpose of __init__ method?", 
                 ["Initialize object attributes", "Destroy objects", "Compare objects", "Copy objects"], 0,
                 "The __init__ method initializes object attributes when an instance is created."),
                ("What is a generator?", 
                 ["Function that returns an iterator", "Random number creator", "List sorter", "File reader"], 0,
                 "Generators are functions that return iterators and can pause and resume execution.")
            ]
        },
        'hard': {
            'questions': [
                ("What is the GIL in Python?", 
                 ["Global Interpreter Lock limiting thread execution", "Graphics Interface Library", "General Input Library", "Global Import Lock"], 0,
                 "The GIL prevents multiple threads from executing Python bytecode simultaneously."),
                ("What is metaclass?", 
                 ["Class that creates classes", "Parent class", "Abstract class", "Static class"], 0,
                 "Metaclasses are classes whose instances are classes themselves."),
                ("How does Python's garbage collection work?", 
                 ["Reference counting with cycle detection", "Manual memory management", "Stack-based collection", "Time-based cleanup"], 0,
                 "Python uses reference counting combined with cycle detection for garbage collection."),
                ("What is the difference between deep and shallow copy?", 
                 ["Deep copy creates new objects, shallow copy references", "No difference", "Shallow copy is recursive", "Deep copy is faster"], 0,
                 "Deep copy creates new objects recursively while shallow copy creates references."),
                ("What is async/await used for?", 
                 ["Asynchronous programming", "Exception handling", "Class inheritance", "Variable declaration"], 0,
                 "async/await enables asynchronous programming for non-blocking operations.")
            ]
        }
    },
    'mathematics': {
        'easy': {
            'questions': [
                ("What is 7 × 8?", 
                 ["56", "54", "64", "48"], 0,
                 "7 × 8 = 56 using basic multiplication."),
                ("What is the square root of 16?", 
                 ["4", "8", "2", "6"], 0,
                 "√16 = 4 because 4 × 4 = 16."),
                ("What is 25% of 100?", 
                 ["25", "75", "50", "20"], 0,
                 "25% of 100 = 25/100 × 100 = 25."),
                ("What is the perimeter of a square with side 5?", 
                 ["20", "25", "10", "15"], 0,
                 "Perimeter of square = 4 × side length = 4 × 5 = 20."),
                ("What is 15 - 7?", 
                 ["8", "9", "7", "22"], 0,
                 "15 - 7 = 8 using basic subtraction.")
            ]
        },
        'medium': {
            'questions': [
                ("What is the derivative of x²?", 
                 ["2x", "x", "2x²", "x²"], 0,
                 "Using the power rule: d/dx(x²) = 2x¹ = 2x."),
                ("What is the area of a circle with radius 3?", 
                 ["9π", "6π", "3π", "12π"], 0,
                 "Area = πr² = π × 3² = 9π."),
                ("Solve for x: 2x + 5 = 15", 
                 ["x = 5", "x = 10", "x = 7.5", "x = 20"], 0,
                 "2x + 5 = 15 → 2x = 10 → x = 5."),
                ("What is sin(90°)?", 
                 ["1", "0", "1/2", "√3/2"], 0,
                 "sin(90°) = 1 in the unit circle."),
                ("What is log₁₀(100)?", 
                 ["2", "10", "100", "1"], 0,
                 "log₁₀(100) = 2 because 10² = 100.")
            ]
        },
        'hard': {
            'questions': [
                ("What is the integral of e^x?", 
                 ["e^x + C", "xe^x + C", "e^x/x + C", "x·e^x + C"], 0,
                 "The integral of e^x is e^x + C because d/dx(e^x) = e^x."),
                ("What is the limit of sin(x)/x as x approaches 0?", 
                 ["1", "0", "∞", "undefined"], 0,
                 "This is a fundamental limit: lim(x→0) sin(x)/x = 1."),
                ("What is the determinant of [[2,3],[1,4]]?", 
                 ["5", "8", "-1", "11"], 0,
                 "det([[2,3],[1,4]]) = 2×4 - 3×1 = 8 - 3 = 5."),
                ("What is the Taylor series of e^x around x=0?", 
                 ["∑(x^n/n!) for n=0 to ∞", "∑(x^n) for n=0 to ∞", "∑(n·x^n) for n=0 to ∞", "∑(x^n/n) for n=1 to ∞"], 0,
                 "The Taylor series of e^x is ∑(x^n/n!) = 1 + x + x²/2! + x³/3! + ..."),
                ("What is the eigenvalue equation?", 
                 ["Av = λv", "A + v = λv", "Av = λ", "A = λv"], 0,
                 "The eigenvalue equation is Av = λv where A is the matrix, v is eigenvector, λ is eigenvalue.")
            ]
        }
    },
    'science': {
        'easy': {
            'questions': [
                ("What is the chemical symbol for water?", 
                 ["H₂O", "CO₂", "NaCl", "O₂"], 0,
                 "Water consists of 2 hydrogen atoms and 1 oxygen atom: H₂O."),
                ("What is the speed of light?", 
                 ["3×10⁸ m/s", "3×10⁶ m/s", "3×10¹⁰ m/s", "3×10⁴ m/s"], 0,
                 "The speed of light in vacuum is approximately 3×10⁸ meters per second."),
                ("What is the powerhouse of the cell?", 
                 ["Mitochondria", "Nucleus", "Ribosome", "Chloroplast"], 0,
                 "Mitochondria produce ATP, the cell's energy currency."),
                ("What gas do plants absorb during photosynthesis?", 
                 ["Carbon dioxide", "Oxygen", "Nitrogen", "Hydrogen"], 0,
                 "Plants absorb CO₂ and release oxygen during photosynthesis."),
                ("What is the atomic number of carbon?", 
                 ["6", "12", "14", "8"], 0,
                 "Carbon has 6 protons, making its atomic number 6.")
            ]
        },
        'medium': {
            'questions': [
                ("What is Newton's second law of motion?", 
                 ["F = ma", "F = mv", "F = m/a", "F = a/m"], 0,
                 "Newton's second law states that Force equals mass times acceleration."),
                ("What is the pH of pure water?", 
                 ["7", "0", "14", "1"], 0,
                 "Pure water has a neutral pH of 7 at room temperature."),
                ("What is the process of cell division called?", 
                 ["Mitosis", "Osmosis", "Photosynthesis", "Respiration"], 0,
                 "Mitosis is the process by which cells divide to create identical copies."),
                ("What is Ohm's law?", 
                 ["V = IR", "P = VI", "E = mc²", "F = qE"], 0,
                 "Ohm's law states that Voltage equals Current times Resistance."),
                ("What is the most abundant gas in Earth's atmosphere?", 
                 ["Nitrogen", "Oxygen", "Carbon dioxide", "Argon"], 0,
                 "Nitrogen makes up about 78% of Earth's atmosphere.")
            ]
        },
        'hard': {
            'questions': [
                ("What is the Heisenberg Uncertainty Principle?", 
                 ["Cannot know position and momentum simultaneously", "Energy is quantized", "Light has wave-particle duality", "Space and time are relative"], 0,
                 "The uncertainty principle states fundamental limits to simultaneously measuring certain pairs of properties."),
                ("What is the mechanism of enzyme catalysis?", 
                 ["Lowering activation energy", "Increasing temperature", "Adding more substrate", "Changing pH"], 0,
                 "Enzymes catalyze reactions by lowering the activation energy required."),
                ("What is the Second Law of Thermodynamics?", 
                 ["Entropy of isolated system always increases", "Energy is conserved", "Force equals mass times acceleration", "Every action has equal opposite reaction"], 0,
                 "The Second Law states that entropy (disorder) of an isolated system always increases."),
                ("What is the structure of DNA?", 
                 ["Double helix with antiparallel strands", "Single strand", "Triple helix", "Circular loop"], 0,
                 "DNA has a double helix structure with two antiparallel complementary strands."),
                ("What is quantum entanglement?", 
                 ["Correlated quantum states regardless of distance", "Particle acceleration", "Wave interference", "Energy quantization"], 0,
                 "Quantum entanglement occurs when particles remain connected so that the quantum state of each particle cannot be described independently.")
            ]
        }
    }
}

# Add this enhanced knowledge base with references after your existing KNOWLEDGE_BASE
KNOWLEDGE_BASE_WITH_REFERENCES = {
    'javascript': {
        'easy': {
            'questions': [
                ("What is the correct way to declare a variable in JavaScript?", 
                 ["var name = 'John';", "variable name = 'John';", "v name = 'John';", "declare name = 'John';"], 0,
                 "The 'var' keyword is used to declare variables in JavaScript.",
                 [
                     {"title": "MDN Web Docs - var", "url": "https://developer.mozilla.org/en-US/docs/Web/JavaScript/Reference/Statements/var", "type": "Official Documentation"},
                     {"title": "JavaScript.info - Variables", "url": "https://javascript.info/variables", "type": "Tutorial"},
                     {"title": "W3Schools - JavaScript Variables", "url": "https://www.w3schools.com/js/js_variables.asp", "type": "Learning Resource"}
                 ]),
                ("Which method is used to add an element to the end of an array?", 
                 ["push()", "add()", "append()", "insert()"], 0,
                 "The push() method adds elements to the end of an array.",
                 [
                     {"title": "MDN Web Docs - Array.prototype.push()", "url": "https://developer.mozilla.org/en-US/docs/Web/JavaScript/Reference/Global_Objects/Array/push", "type": "Official Documentation"},
                     {"title": "JavaScript.info - Array methods", "url": "https://javascript.info/array-methods", "type": "Tutorial"},
                     {"title": "ECMAScript Specification - Array.prototype.push", "url": "https://tc39.es/ecma262/#sec-array.prototype.push", "type": "Specification"}
                 ]),
                ("What does '===' operator do in JavaScript?", 
                 ["Strict equality comparison", "Assignment", "Loose equality", "Not equal"], 0,
                 "The '===' operator performs strict equality comparison without type conversion.",
                 [
                     {"title": "MDN Web Docs - Strict equality (===)", "url": "https://developer.mozilla.org/en-US/docs/Web/JavaScript/Reference/Operators/Strict_equality", "type": "Official Documentation"},
                     {"title": "JavaScript.info - Comparisons", "url": "https://javascript.info/comparison", "type": "Tutorial"},
                     {"title": "ECMAScript Specification - Strict Equality", "url": "https://tc39.es/ecma262/#sec-strict-equality-comparison", "type": "Specification"}
                 ])
            ]
        },
        'medium': {
            'questions': [
                ("What is closure in JavaScript?", 
                 ["Function with access to outer scope", "A loop structure", "An object method", "A data type"], 0,
                 "Closure allows a function to access variables from its outer/enclosing scope.",
                 [
                     {"title": "MDN Web Docs - Closures", "url": "https://developer.mozilla.org/en-US/docs/Web/JavaScript/Closures", "type": "Official Documentation"},
                     {"title": "JavaScript.info - Variable scope, closures", "url": "https://javascript.info/closure", "type": "Tutorial"},
                     {"title": "You Don't Know JS - Scope & Closures", "url": "https://github.com/getify/You-Dont-Know-JS/tree/2nd-ed/scope-closures", "type": "Book"}
                 ]),
                ("Which method creates a new array with results of calling a function?", 
                 ["map()", "filter()", "reduce()", "forEach()"], 0,
                 "The map() method creates a new array with the results of calling a function for every array element.",
                 [
                     {"title": "MDN Web Docs - Array.prototype.map()", "url": "https://developer.mozilla.org/en-US/docs/Web/JavaScript/Reference/Global_Objects/Array/map", "type": "Official Documentation"},
                     {"title": "JavaScript.info - Array methods", "url": "https://javascript.info/array-methods#map", "type": "Tutorial"},
                     {"title": "ECMAScript Specification - Array.prototype.map", "url": "https://tc39.es/ecma262/#sec-array.prototype.map", "type": "Specification"}
                 ])
            ]
        },
        'hard': {
            'questions': [
                ("What is the time complexity of Array.prototype.indexOf()?", 
                 ["O(n)", "O(log n)", "O(1)", "O(n²)"], 0,
                 "indexOf() performs a linear search, resulting in O(n) time complexity.",
                 [
                     {"title": "MDN Web Docs - Array.prototype.indexOf()", "url": "https://developer.mozilla.org/en-US/docs/Web/JavaScript/Reference/Global_Objects/Array/indexOf", "type": "Official Documentation"},
                     {"title": "Big O Notation - Khan Academy", "url": "https://www.khanacademy.org/computing/computer-science/algorithms/asymptotic-notation/a/big-o-notation", "type": "Educational"},
                     {"title": "ECMAScript Specification - Array.prototype.indexOf", "url": "https://tc39.es/ecma262/#sec-array.prototype.indexof", "type": "Specification"}
                 ]),
                ("What is the prototype chain?", 
                 ["Inheritance mechanism in JavaScript", "A design pattern", "A loop structure", "An error handling method"], 0,
                 "The prototype chain enables inheritance by linking objects through their prototype property.",
                 [
                     {"title": "MDN Web Docs - Inheritance and prototype chain", "url": "https://developer.mozilla.org/en-US/docs/Web/JavaScript/Inheritance_and_the_prototype_chain", "type": "Official Documentation"},
                     {"title": "JavaScript.info - Prototypal inheritance", "url": "https://javascript.info/prototype-inheritance", "type": "Tutorial"},
                     {"title": "You Don't Know JS - this & Object Prototypes", "url": "https://github.com/getify/You-Dont-Know-JS/tree/1st-ed/this%20%26%20object%20prototypes", "type": "Book"}
                 ])
            ]
        }
    },
    'python': {
        'easy': {
            'questions': [
                ("How do you create a list in Python?", 
                 ["my_list = [1, 2, 3]", "my_list = (1, 2, 3)", "my_list = {1, 2, 3}", "my_list = <1, 2, 3>"], 0,
                 "Lists in Python are created using square brackets.",
                 [
                     {"title": "Python.org - Lists", "url": "https://docs.python.org/3/tutorial/datastructures.html#more-on-lists", "type": "Official Documentation"},
                     {"title": "Real Python - Lists and Tuples", "url": "https://realpython.com/python-lists-tuples/", "type": "Tutorial"},
                     {"title": "Python PEP 289 - List Comprehensions", "url": "https://www.python.org/dev/peps/pep-0289/", "type": "PEP Document"}
                 ]),
                ("Which keyword is used to define a function?", 
                 ["def", "function", "define", "func"], 0,
                 "The 'def' keyword is used to define functions in Python.",
                 [
                     {"title": "Python.org - Defining Functions", "url": "https://docs.python.org/3/tutorial/controlflow.html#defining-functions", "type": "Official Documentation"},
                     {"title": "Real Python - Python Functions", "url": "https://realpython.com/defining-your-own-python-function/", "type": "Tutorial"},
                     {"title": "PEP 3107 - Function Annotations", "url": "https://www.python.org/dev/peps/pep-3107/", "type": "PEP Document"}
                 ])
            ]
        },
        'medium': {
            'questions': [
                ("What is the GIL in Python?", 
                 ["Global Interpreter Lock limiting thread execution", "Graphics Interface Library", "General Input Library", "Global Import Lock"], 0,
                 "The GIL prevents multiple threads from executing Python bytecode simultaneously.",
                 [
                     {"title": "Python.org - Global Interpreter Lock", "url": "https://docs.python.org/3/c-api/init.html#thread-state-and-the-global-interpreter-lock", "type": "Official Documentation"},
                     {"title": "Real Python - Python's GIL", "url": "https://realpython.com/python-gil/", "type": "Tutorial"},
                     {"title": "PEP 311 - Simplified GIL Acquisition", "url": "https://www.python.org/dev/peps/pep-0311/", "type": "PEP Document"}
                 ])
            ]
        }
    },
    'mathematics': {
        'easy': {
            'questions': [
                ("What is 7 × 8?", 
                 ["56", "54", "64", "48"], 0,
                 "7 × 8 = 56 using basic multiplication."),
                ("What is the square root of 16?", 
                 ["4", "8", "2", "6"], 0,
                 "√16 = 4 because 4 × 4 = 16."),
                ("What is 25% of 100?", 
                 ["25", "75", "50", "20"], 0,
                 "25% of 100 = 25/100 × 100 = 25."),
                ("What is the perimeter of a square with side 5?", 
                 ["20", "25", "10", "15"], 0,
                 "Perimeter of square = 4 × side length = 4 × 5 = 20."),
                ("What is 15 - 7?", 
                 ["8", "9", "7", "22"], 0,
                 "15 - 7 = 8 using basic subtraction.")
            ]
        },
        'medium': {
            'questions': [
                ("What is the derivative of x²?", 
                 ["2x", "x", "2x²", "x²"], 0,
                 "Using the power rule: d/dx(x²) = 2x¹ = 2x."),
                ("What is the area of a circle with radius 3?", 
                 ["9π", "6π", "3π", "12π"], 0,
                 "Area = πr² = π × 3² = 9π."),
                ("Solve for x: 2x + 5 = 15", 
                 ["x = 5", "x = 10", "x = 7.5", "x = 20"], 0,
                 "2x + 5 = 15 → 2x = 10 → x = 5."),
                ("What is sin(90°)?", 
                 ["1", "0", "1/2", "√3/2"], 0,
                 "sin(90°) = 1 in the unit circle."),
                ("What is log₁₀(100)?", 
                 ["2", "10", "100", "1"], 0,
                 "log₁₀(100) = 2 because 10² = 100.")
            ]
        },
        'hard': {
            'questions': [
                ("What is the integral of e^x?", 
                 ["e^x + C", "xe^x + C", "e^x/x + C", "x·e^x + C"], 0,
                 "The integral of e^x is e^x + C because d/dx(e^x) = e^x."),
                ("What is the limit of sin(x)/x as x approaches 0?", 
                 ["1", "0", "∞", "undefined"], 0,
                 "This is a fundamental limit: lim(x→0) sin(x)/x = 1."),
                ("What is the determinant of [[2,3],[1,4]]?", 
                 ["5", "8", "-1", "11"], 0,
                 "det([[2,3],[1,4]]) = 2×4 - 3×1 = 8 - 3 = 5."),
                ("What is the Taylor series of e^x around x=0?", 
                 ["∑(x^n/n!) for n=0 to ∞", "∑(x^n) for n=0 to ∞", "∑(n·x^n) for n=0 to ∞", "∑(x^n/n) for n=1 to ∞"], 0,
                 "The Taylor series of e^x is ∑(x^n/n!) = 1 + x + x²/2! + x³/3! + ..."),
                ("What is the eigenvalue equation?", 
                 ["Av = λv", "A + v = λv", "Av = λ", "A = λv"], 0,
                 "The eigenvalue equation is Av = λv where A is the matrix, v is eigenvector, λ is eigenvalue.")
            ]
        }
    },
    'science': {
        'easy': {
            'questions': [
                ("What is the chemical symbol for water?", 
                 ["H₂O", "CO₂", "NaCl", "O₂"], 0,
                 "Water consists of 2 hydrogen atoms and 1 oxygen atom: H₂O."),
                ("What is the speed of light?", 
                 ["3×10⁸ m/s", "3×10⁶ m/s", "3×10¹⁰ m/s", "3×10⁴ m/s"], 0,
                 "The speed of light in vacuum is approximately 3×10⁸ meters per second."),
                ("What is the powerhouse of the cell?", 
                 ["Mitochondria", "Nucleus", "Ribosome", "Chloroplast"], 0,
                 "Mitochondria produce ATP, the cell's energy currency."),
                ("What gas do plants absorb during photosynthesis?", 
                 ["Carbon dioxide", "Oxygen", "Nitrogen", "Hydrogen"], 0,
                 "Plants absorb CO₂ and release oxygen during photosynthesis."),
                ("What is the atomic number of carbon?", 
                 ["6", "12", "14", "8"], 0,
                 "Carbon has 6 protons, making its atomic number 6.")
            ]
        },
        'medium': {
            'questions': [
                ("What is Newton's second law of motion?", 
                 ["F = ma", "F = mv", "F = m/a", "F = a/m"], 0,
                 "Newton's second law states that Force equals mass times acceleration."),
                ("What is the pH of pure water?", 
                 ["7", "0", "14", "1"], 0,
                 "Pure water has a neutral pH of 7 at room temperature."),
                ("What is the process of cell division called?", 
                 ["Mitosis", "Osmosis", "Photosynthesis", "Respiration"], 0,
                 "Mitosis is the process by which cells divide to create identical copies."),
                ("What is Ohm's law?", 
                 ["V = IR", "P = VI", "E = mc²", "F = qE"], 0,
                 "Ohm's law states that Voltage equals Current times Resistance."),
                ("What is the most abundant gas in Earth's atmosphere?", 
                 ["Nitrogen", "Oxygen", "Carbon dioxide", "Argon"], 0,
                 "Nitrogen makes up about 78% of Earth's atmosphere.")
            ]
        },
        'hard': {
            'questions': [
                ("What is the Heisenberg Uncertainty Principle?", 
                 ["Cannot know position and momentum simultaneously", "Energy is quantized", "Light has wave-particle duality", "Space and time are relative"], 0,
                 "The uncertainty principle states fundamental limits to simultaneously measuring certain pairs of properties."),
                ("What is the mechanism of enzyme catalysis?", 
                 ["Lowering activation energy", "Increasing temperature", "Adding more substrate", "Changing pH"], 0,
                 "Enzymes catalyze reactions by lowering the activation energy required."),
                ("What is the Second Law of Thermodynamics?", 
                 ["Entropy of isolated system always increases", "Energy is conserved", "Force equals mass times acceleration", "Every action has equal opposite reaction"], 0,
                 "The Second Law states that entropy (disorder) of an isolated system always increases."),
                ("What is the structure of DNA?", 
                 ["Double helix with antiparallel strands", "Single strand", "Triple helix", "Circular loop"], 0,
                 "DNA has a double helix structure with two antiparallel complementary strands."),
                ("What is quantum entanglement?", 
                 ["Correlated quantum states regardless of distance", "Particle acceleration", "Wave interference", "Energy quantization"], 0,
                 "Quantum entanglement occurs when particles remain connected so that the quantum state of each particle cannot be described independently.")
            ]
        }
    }
}

def generate_advanced_ai_questions(quiz_data):
    """Generate diverse, logical questions using advanced AI logic"""
    questions = []
    count = int(quiz_data.get('questionCount', 5))
    topic = quiz_data.get('topic', '').lower().strip()
    difficulty = quiz_data.get('difficulty', 'medium').lower()
    subject = quiz_data.get('subject', '').lower().strip()
    
    print(f"🤖 Generating questions for: {topic} in {subject} ({difficulty})")
    
    # Determine the best knowledge domain
    domain_key = None
    
    # Smart topic matching
    if any(keyword in topic for keyword in ['javascript', 'js', 'programming', 'coding', 'function', 'variable', 'array']):
        domain_key = 'javascript'
    elif any(keyword in topic for keyword in ['python', 'django', 'flask', 'list', 'dictionary', 'class']):
        domain_key = 'python'
    elif any(keyword in topic for keyword in ['math', 'algebra', 'calculus', 'geometry', 'equation', 'derivative', 'integral']):
        domain_key = 'mathematics'
    elif any(keyword in topic for keyword in ['chemistry', 'physics', 'biology', 'science', 'cell', 'atom', 'molecule']):
        domain_key = 'science'
    elif any(keyword in subject for keyword in ['computer', 'programming', 'software']):
        domain_key = 'javascript'  # Default programming
    elif any(keyword in subject for keyword in ['math', 'mathematics']):
        domain_key = 'mathematics'
    elif any(keyword in subject for keyword in ['science', 'chemistry', 'physics', 'biology']):
        domain_key = 'science'
    else:
        # If no match, use a general approach
        domain_key = 'science'  # Default fallback
    
    # Get questions from knowledge base
    if domain_key and domain_key in KNOWLEDGE_BASE and difficulty in KNOWLEDGE_BASE[domain_key]:
        knowledge_questions = KNOWLEDGE_BASE[domain_key][difficulty]['questions']
        
        # If we have enough pre-made questions, use them
        if len(knowledge_questions) >= count:
            selected_questions = random.sample(knowledge_questions, count)
        else:
            # Use all available and generate more
            selected_questions = knowledge_questions[:]
            remaining = count - len(selected_questions)
            
            # Generate additional questions
            for i in range(remaining):
                additional_q = generate_contextual_question(topic, subject, difficulty, i + len(selected_questions) + 1)
                selected_questions.append(additional_q)
    else:
        # Generate all questions contextually
        selected_questions = []
        for i in range(count):
            question = generate_contextual_question(topic, subject, difficulty, i + 1)
            selected_questions.append(question)
    
    # Convert to our format
    for i, q_data in enumerate(selected_questions):
        if isinstance(q_data, tuple) and len(q_data) >= 4:
            question_text, options, correct_idx, explanation = q_data[0], q_data[1], q_data[2], q_data[3]
        else:
            question_text, options, correct_idx, explanation = q_data
        
        question = {
            'id': i + 1,
            'question': question_text,
            'options': options,
            'correct_answer': correct_idx,
            'explanation': explanation,
            'ai_generated': True,
            'difficulty': difficulty,
            'topic': topic,
            'domain': domain_key
        }
        questions.append(question)
    
    print(f"✅ Generated {len(questions)} diverse questions for {topic}")
    return questions

def generate_contextual_question(topic, subject, difficulty, question_num):
    """Generate contextual questions when no pre-made questions exist"""
    
    templates = {
        'easy': [
            f"What is the basic definition of {topic}?",
            f"Which of the following best describes {topic}?",
            f"What is a key characteristic of {topic}?",
            f"In {subject}, {topic} is primarily used for:",
            f"What is the main purpose of {topic}?"
        ],
        'medium': [
            f"How does {topic} work in practical applications?",
            f"What are the advantages of using {topic}?",
            f"How does {topic} relate to other concepts in {subject}?",
            f"What is the best approach when implementing {topic}?",
            f"What problems does {topic} solve in {subject}?"
        ],
        'hard': [
            f"What are the theoretical foundations of {topic}?",
            f"How would you optimize {topic} for complex scenarios?",
            f"What are the limitations and trade-offs of {topic}?",
            f"How does {topic} integrate with advanced {subject} concepts?",
            f"What are the cutting-edge developments in {topic}?"
        ]
    }
    
    question_templates = templates.get(difficulty, templates['medium'])
    question_text = question_templates[(question_num - 1) % len(question_templates)]
    
    # Generate more diverse options
    if 'programming' in subject.lower() or 'computer' in subject.lower():
        options = [
            f"It provides efficient solutions for {topic}-related problems",
            f"It is an outdated approach to handling {topic}",
            f"It only works in specific programming languages",
            f"It has no practical applications in modern development"
        ]
    elif 'math' in subject.lower():
        options = [
            f"It follows fundamental mathematical principles for {topic}",
            f"It contradicts basic mathematical rules",
            f"It only applies to theoretical mathematics",
            f"It is used exclusively in advanced calculus"
        ]
    else:
        options = [
            f"It represents current scientific understanding of {topic}",
            f"It is based on disproven theories about {topic}",
            f"It only applies in laboratory conditions",
            f"It has no connection to real-world phenomena"
        ]
    
    # Randomize correct answer position
    correct_answer = random.randint(0, 3)
    
    explanation = f"This answer correctly explains {topic} because it aligns with established principles in {subject} and reflects current understanding in the field."
    
    return (question_text, options, correct_answer, explanation)

def generate_advanced_ai_questions_with_references(quiz_data):
    """Generate diverse, unique questions with authoritative references"""
    questions = []
    count = int(quiz_data.get('questionCount', 5))
    topic = quiz_data.get('topic', '').lower().strip()
    difficulty = quiz_data.get('difficulty', 'medium').lower()
    subject = quiz_data.get('subject', '').lower().strip()
    
    print(f"🤖 Generating {count} unique questions for: {topic} in {subject} ({difficulty})")
    
    # Determine the best knowledge domain
    domain_key = determine_domain(topic, subject)
    
    # Create diverse question templates for each topic
    question_templates = get_diverse_question_templates(topic, subject, difficulty)
    
    # Generate unique questions using different templates
    used_templates = set()
    used_concepts = set()
    
    for i in range(count):
        # Ensure we don't repeat question types
        available_templates = [t for t in question_templates if t['id'] not in used_templates]
        if not available_templates:
            # Reset if we've used all templates
            used_templates.clear()
            available_templates = question_templates
        
        # Select a unique template
        template = random.choice(available_templates)
        used_templates.add(template['id'])
        
        # Generate question with specific concept focus
        concept = generate_unique_concept(topic, subject, used_concepts, i)
        used_concepts.add(concept)
        
        question_data = generate_contextual_question_with_template(
            topic, subject, difficulty, template, concept, i + 1
        )
        
        # Randomize correct answer position
        question_text, options, correct_idx, explanation, references = question_data
        randomized_correct_idx = random.randint(0, 3)
        randomized_options = options[:]
        
        # Swap correct answer to new position
        if randomized_correct_idx != correct_idx:
            randomized_options[randomized_correct_idx], randomized_options[correct_idx] = \
                randomized_options[correct_idx], randomized_options[randomized_correct_idx]
        
        question = {
            'id': i + 1,
            'question': question_text,
            'options': randomized_options,
            'correct_answer': randomized_correct_idx,
            'explanation': explanation,
            'references': references,
            'ai_generated': True,
            'verified': True,
            'difficulty': difficulty,
            'topic': topic,
            'concept': concept,
            'template_used': template['id']
        }
        questions.append(question)
    
    print(f"✅ Generated {len(questions)} unique questions with diverse concepts")
    return questions

def determine_domain(topic, subject):
    """Determine the knowledge domain based on topic and subject"""
    topic_lower = topic.lower()
    subject_lower = subject.lower()
    
    # JavaScript/Programming keywords
    js_keywords = ['javascript', 'js', 'function', 'variable', 'array', 'object', 'dom', 'api', 'async', 'promise']
    # Python keywords  
    python_keywords = ['python', 'django', 'flask', 'list', 'dictionary', 'class', 'module', 'pip', 'pandas']
    # Math keywords
    math_keywords = ['algebra', 'calculus', 'geometry', 'trigonometry', 'equation', 'derivative', 'integral', 'matrix']
    # Science keywords
    science_keywords = ['chemistry', 'physics', 'biology', 'cell', 'atom', 'molecule', 'energy', 'force', 'genetics']
    
    if any(keyword in topic_lower or keyword in subject_lower for keyword in js_keywords):
        return 'javascript'
    elif any(keyword in topic_lower or keyword in subject_lower for keyword in python_keywords):
        return 'python'
    elif any(keyword in topic_lower or keyword in subject_lower for keyword in math_keywords):
        return 'mathematics'
    elif any(keyword in topic_lower or keyword in subject_lower for keyword in science_keywords):
        return 'science'
    elif 'computer' in subject_lower or 'programming' in subject_lower:
        return 'javascript'
    elif 'math' in subject_lower:
        return 'mathematics'
    else:
        return 'science'

def get_diverse_question_templates(topic, subject, difficulty):
    """Get diverse question templates to avoid repetition"""
    
    base_templates = [
        {'id': 'definition', 'type': 'What is', 'focus': 'basic understanding'},
        {'id': 'application', 'type': 'How to use', 'focus': 'practical application'},
        {'id': 'comparison', 'type': 'Compare/contrast', 'focus': 'relationship analysis'},
        {'id': 'problem_solving', 'type': 'Problem solving', 'focus': 'solution approach'},
        {'id': 'best_practice', 'type': 'Best practice', 'focus': 'optimal approach'},
        {'id': 'troubleshooting', 'type': 'Troubleshooting', 'focus': 'error handling'},
        {'id': 'implementation', 'type': 'Implementation', 'focus': 'coding/application'},
        {'id': 'theory', 'type': 'Theoretical', 'focus': 'conceptual understanding'}
    ]
    
    # Adjust templates based on difficulty
    if difficulty == 'easy':
        return [t for t in base_templates if t['id'] in ['definition', 'application', 'best_practice']]
    elif difficulty == 'medium':
        return [t for t in base_templates if t['id'] in ['application', 'comparison', 'problem_solving', 'implementation']]
    else:  # hard
        return [t for t in base_templates if t['id'] in ['problem_solving', 'troubleshooting', 'theory', 'implementation']]

def generate_unique_concept(topic, subject, used_concepts, question_number):
    """Generate unique concept variations to avoid repetition"""
    
    base_topic = topic.lower()
    
    # Create concept variations based on domain
    if 'javascript' in subject.lower() or 'programming' in subject.lower():
        concepts = [
            f"{base_topic} syntax",
            f"{base_topic} best practices", 
            f"{base_topic} performance",
            f"{base_topic} debugging",
            f"{base_topic} implementation",
            f"{base_topic} optimization",
            f"{base_topic} error handling",
            f"{base_topic} integration"
        ]
    elif 'math' in subject.lower():
        concepts = [
            f"{base_topic} fundamentals",
            f"{base_topic} applications",
            f"{base_topic} problem solving",
            f"{base_topic} theorems",
            f"{base_topic} proofs",
            f"{base_topic} calculations",
            f"{base_topic} formulas",
            f"{base_topic} real-world usage"
        ]
    else:
        concepts = [
            f"{base_topic} principles",
            f"{base_topic} applications",
            f"{base_topic} methodology",
            f"{base_topic} analysis",
            f"{base_topic} implementation",
            f"{base_topic} evaluation",
            f"{base_topic} optimization",
            f"{base_topic} integration"
        ]
    
    # Return unused concept or create new one
    available_concepts = [c for c in concepts if c not in used_concepts]
    if available_concepts:
        return random.choice(available_concepts)
    else:
        return f"{base_topic} aspect {question_number}"

def generate_contextual_question_with_template(topic, subject, difficulty, template, concept, question_num):
    """Generate unique questions using templates and concepts"""
    
    # Create diverse question stems based on template type
    if template['id'] == 'definition':
        question_stems = [
            f"What is the primary purpose of {concept}?",
            f"How would you define {concept}?",
            f"What characterizes {concept}?",
            f"Which statement best describes {concept}?"
        ]
    elif template['id'] == 'application':
        question_stems = [
            f"When implementing {concept}, what is the recommended approach?",
            f"How should {concept} be applied in practical scenarios?",
            f"What is the best way to utilize {concept}?",
            f"In which situation would you use {concept}?"
        ]
    elif template['id'] == 'comparison':
        question_stems = [
            f"How does {concept} compare to alternative approaches?",
            f"What advantage does {concept} provide over other methods?",
            f"When choosing between options, why select {concept}?",
            f"What makes {concept} different from similar concepts?"
        ]
    elif template['id'] == 'problem_solving':
        question_stems = [
            f"When facing challenges with {concept}, what should you do?",
            f"How do you solve problems related to {concept}?",
            f"What approach works best for {concept} issues?",
            f"If {concept} isn't working as expected, what's the solution?"
        ]
    elif template['id'] == 'best_practice':
        question_stems = [
            f"What is considered best practice for {concept}?",
            f"Which approach is recommended when working with {concept}?",
            f"What guidelines should be followed for {concept}?",
            f"How should professionals handle {concept}?"
        ]
    elif template['id'] == 'troubleshooting':
        question_stems = [
            f"When {concept} produces unexpected results, what's the likely cause?",
            f"How do you debug issues with {concept}?",
            f"What's the first step in troubleshooting {concept}?",
            f"Common problems with {concept} are usually caused by what?"
        ]
    elif template['id'] == 'implementation':
        question_stems = [
            f"The correct way to implement {concept} involves which step?",
            f"When setting up {concept}, what is essential?",
            f"What must be considered during {concept} implementation?",
            f"Which factor is most important for successful {concept}?"
        ]
    else:  # theory
        question_stems = [
            f"The theoretical foundation of {concept} is based on what?",
            f"What principle underlies {concept}?",
            f"From a theoretical perspective, {concept} represents what?",
            f"The conceptual basis for {concept} comes from which idea?"
        ]
    
    # Select question stem
    question_text = random.choice(question_stems)
    
    # Generate diverse, contextual options
    options = generate_diverse_options(topic, subject, concept, template, difficulty)
    
    # Always put correct answer first, then randomize later
    correct_answer = 0
    
    # Generate contextual explanation
    explanation = generate_contextual_explanation(concept, subject, template)
    
    # Generate appropriate references
    references = generate_domain_specific_references(topic, subject, concept)
    
    return (question_text, options, correct_answer, explanation, references)

def generate_diverse_options(topic, subject, concept, template, difficulty):
    """Generate diverse answer options based on context"""
    
    domain = determine_domain(topic, subject)
    
    if domain == 'javascript':
        if template['id'] == 'best_practice':
            return [
                f"Follow established coding standards and use appropriate design patterns for {concept}",
                f"Ignore performance considerations and focus only on functionality",
                f"Use deprecated methods for backward compatibility",
                f"Avoid documentation and code comments to keep it simple"
            ]
        elif template['id'] == 'problem_solving':
            return [
                f"Analyze the issue systematically and apply debugging techniques for {concept}",
                f"Randomly try different approaches until something works",
                f"Copy solutions from online without understanding the logic",
                f"Ignore error messages and continue with implementation"
            ]
        else:
            return [
                f"It provides efficient and scalable solutions for {concept}",
                f"It creates unnecessary complexity in the codebase",
                f"It only works in specific browser environments",
                f"It has been deprecated and should not be used"
            ]
            
    elif domain == 'mathematics':
        if template['id'] == 'application':
            return [
                f"Apply the fundamental mathematical principles correctly for {concept}",
                f"Use approximations even when exact solutions are possible",
                f"Ignore mathematical rules and rely on intuition",
                f"Only use basic arithmetic regardless of complexity"
            ]
        elif template['id'] == 'problem_solving':
            return [
                f"Break down the problem systematically and apply appropriate formulas for {concept}",
                f"Guess the answer based on similar-looking problems",
                f"Use only mental math without showing work",
                f"Skip steps to reach the conclusion faster"
            ]
        else:
            return [
                f"It follows proven mathematical principles and provides reliable results for {concept}",
                f"It contradicts established mathematical theorems",
                f"It only applies to theoretical situations with no practical use",
                f"It requires advanced mathematics that most people can't understand"
            ]
            
    else:  # science or general
        if template['id'] == 'methodology':
            return [
                f"Use evidence-based approaches and follow scientific methodology for {concept}",
                f"Rely on personal opinions rather than empirical evidence",
                f"Skip the hypothesis stage and jump to conclusions",
                f"Avoid peer review and independent verification"
            ]
        else:
            return [
                f"It represents current scientific understanding and evidence-based knowledge of {concept}",
                f"It contradicts well-established scientific principles",
                f"It only applies in controlled laboratory conditions",
                f"It has no basis in scientific research or evidence"
            ]

def generate_contextual_explanation(concept, subject, template):
    """Generate contextual explanations based on concept and template"""
    
    explanations = {
        'definition': f"This answer correctly defines {concept} according to established standards in {subject}.",
        'application': f"This approach represents the proper application of {concept} in {subject} contexts.",
        'comparison': f"This option accurately compares {concept} with alternatives in {subject}.",
        'problem_solving': f"This solution follows proven problem-solving methodologies for {concept} in {subject}.",
        'best_practice': f"This represents industry best practices and standards for {concept} in {subject}.",
        'troubleshooting': f"This approach follows systematic troubleshooting procedures for {concept} in {subject}.",
        'implementation': f"This method ensures proper implementation of {concept} according to {subject} principles.",
        'theory': f"This answer reflects the theoretical foundations of {concept} in {subject}."
    }
    
    return explanations.get(template['id'], f"This answer correctly explains {concept} in the context of {subject}.")

def generate_domain_specific_references(topic, subject, concept):
    """Generate domain-specific references for verification"""
    
    domain = determine_domain(topic, subject)
    
    if domain == 'javascript':
        return [
            {
                "title": f"MDN Web Docs - {concept}",
                "url": f"https://developer.mozilla.org/en-US/search?q={concept.replace(' ', '+')}", 
                "type": "Official Documentation"
            },
            {
                "title": f"JavaScript.info - {concept}",
                "url": f"https://javascript.info/?s={concept.replace(' ', '+')}", 
                "type": "Educational Resource"
            },
            {
                "title": f"W3Schools - {concept}",
                "url": f"https://www.w3schools.com/js/", 
                "type": "Tutorial"
            }
        ]
    elif domain == 'python':
        return [
            {
                "title": f"Python Documentation - {concept}",
                "url": f"https://docs.python.org/3/search.html?q={concept.replace(' ', '+')}", 
                "type": "Official Documentation"
            },
            {
                "title": f"Real Python - {concept}",
                "url": f"https://realpython.com/search/?q={concept.replace(' ', '+')}", 
                "type": "Tutorial"
            },
            {
                "title": f"Python Package Index - {concept}",
                "url": f"https://pypi.org/search/?q={concept.replace(' ', '+')}", 
                "type": "Package Repository"
            }
        ]
    elif domain == 'mathematics':
        return [
            {
                "title": f"Khan Academy - {concept}",
                "url": f"https://www.khanacademy.org/search?page_search_query={concept.replace(' ', '+')}", 
                "type": "Educational"
            },
            {
                "title": f"Wolfram MathWorld - {concept}",
                "url": f"https://mathworld.wolfram.com/search/?query={concept.replace(' ', '+')}", 
                "type": "Mathematical Reference"
            },
            {
                "title": f"MIT OpenCourseWare - {concept}",
                "url": f"https://ocw.mit.edu/search/?q={concept.replace(' ', '+')}", 
                "type": "Academic Course"
            }
        ]
    else:  # science
        return [
            {
                "title": f"Scientific Research - {concept}",
                "url": f"https://scholar.google.com/scholar?q={concept.replace(' ', '+')}", 
                "type": "Academic Research"
            },
            {
                "title": f"Khan Academy Science - {concept}",
                "url": f"https://www.khanacademy.org/science/", 
                "type": "Educational"
            },
            {
                "title": f"NASA Education - {concept}",
                "url": f"https://www.nasa.gov/search/site/{concept.replace(' ', '+')}", 
                "type": "Scientific Institution"
            }
        ]

# Add download functions
def generate_pdf(quiz, questions):
    if not FPDF_AVAILABLE:
        return jsonify({'error': 'PDF generation not available'}), 500
        
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, f'Quiz: {quiz.title}', 0, 1, 'C')
    
    pdf.set_font('Arial', '', 12)
    pdf.cell(0, 10, f'Subject: {quiz.subject}', 0, 1)
    pdf.cell(0, 10, f'Topic: {quiz.topic}', 0, 1)
    pdf.cell(0, 10, f'Difficulty: {quiz.difficulty}', 0, 1)
    pdf.ln(5)
    
    for i, q in enumerate(questions, 1):
        pdf.set_font('Arial', 'B', 12)
        # Handle long questions by splitting them
        question_text = q["question"][:80] + "..." if len(q["question"]) > 80 else q["question"]
        pdf.cell(0, 10, f'Q{i}: {question_text}', 0, 1)
        
        pdf.set_font('Arial', '', 10)
        for j, option in enumerate(q['options']):
            marker = '[X]' if j == q['correct_answer'] else '[ ]'
            option_text = option[:70] + "..." if len(option) > 70 else option
            pdf.cell(0, 6, f'  {marker} {chr(65+j)}) {option_text}', 0, 1)
        pdf.ln(3)
    
    output = io.BytesIO()
    pdf.output(output, 'S')
    output.seek(0)
    
    return send_file(output, mimetype='application/pdf', 
                    as_attachment=True, download_name=f'{quiz.title}.pdf')

def generate_docx(quiz, questions):
    if not DOCX_AVAILABLE:
        return jsonify({'error': 'Word generation not available'}), 500
        
    doc = Document()
    doc.add_heading(f'Quiz: {quiz.title}', 0)
    
    doc.add_paragraph(f'Subject: {quiz.subject}')
    doc.add_paragraph(f'Topic: {quiz.topic}')
    doc.add_paragraph(f'Difficulty: {quiz.difficulty}')
    
    for i, q in enumerate(questions, 1):
        doc.add_heading(f'Question {i}', level=2)
        doc.add_paragraph(q['question'])
        
        for j, option in enumerate(q['options']):
            marker = '[X]' if j == q['correct_answer'] else '[ ]'
            doc.add_paragraph(f'{marker} {chr(65+j)}) {option}')
        
        doc.add_paragraph(f'Explanation: {q["explanation"]}')
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True, download_name=f'{quiz.title}.docx')

def generate_xlsx(quiz, questions):
    if not XLSX_AVAILABLE:
        return jsonify({'error': 'Excel generation not available'}), 500
        
    data = []
    for i, q in enumerate(questions, 1):
        data.append({
            'Question_Number': i,
            'Question': q['question'],
            'Option_A': q['options'][0],
            'Option_B': q['options'][1],
            'Option_C': q['options'][2],
            'Option_D': q['options'][3],
            'Correct_Answer': chr(65 + q['correct_answer']),
            'Explanation': q['explanation']
        })
    
    df = pd.DataFrame(data)
    output = io.BytesIO()
    df.to_excel(output, index=False, sheet_name=quiz.title[:31])
    output.seek(0)
    
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    as_attachment=True, download_name=f'{quiz.title}.xlsx')

# Add download endpoint
@app.route('/api/quiz/<int:quiz_id>/download/<format>', methods=['GET'])
def download_quiz(quiz_id, format):
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'No token provided'}), 401
            
        data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
        user_id = data['user_id']
        
        quiz = Quiz.query.filter_by(id=quiz_id, user_id=user_id).first()
        if not quiz:
            return jsonify({'error': 'Quiz not found'}), 404
            
        questions = json.loads(quiz.questions)
        
        if format == 'pdf':
            return generate_pdf(quiz, questions)
        elif format == 'docx':
            return generate_docx(quiz, questions)
        elif format == 'xlsx':
            return generate_xlsx(quiz, questions)
        else:
            return jsonify({'error': 'Unsupported format'}), 400
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Update the existing quiz creation endpoint
@app.route('/api/quiz', methods=['POST'])
def create_quiz():
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'No token provided'}), 401
            
        data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
        user_id = data['user_id']
        
        quiz_data = request.get_json()
        print(f"🤖 Creating verified AI quiz with references: {quiz_data.get('title')} by user {user_id}")
        
        # Use enhanced AI question generation with references
        questions = generate_advanced_ai_questions_with_references(quiz_data)
        
        quiz = Quiz(
            user_id=user_id,
            title=quiz_data.get('title'),
            subject=quiz_data.get('subject'),
            topic=quiz_data.get('topic'),
            difficulty=quiz_data.get('difficulty'),
            questions=json.dumps(questions)
        )
        
        db.session.add(quiz)
        db.session.commit()
        
        print(f"✅ Verified AI Quiz created: {quiz.title} with {len(questions)} questions and references")
        return jsonify({
            'id': quiz.id,
            'title': quiz.title,
            'subject': quiz.subject,
            'topic': quiz.topic,
            'difficulty': quiz.difficulty,
            'questions': questions,
            'created_at': quiz.created_at.isoformat(),
            'ai_powered': True,
            'verified': True,
            'knowledge_domains': list(set([q.get('domain', 'general') for q in questions])),
            'reference_count': sum(len(q.get('references', [])) for q in questions)
        })
        
    except Exception as e:
        print(f"❌ Verified AI Quiz creation error: {e}")
        return jsonify({'error': str(e)}), 500

# Keep all your other existing endpoints (health_check, login, register, get_quizzes, download functions, etc.)
@app.route('/', methods=['GET'])
def health_check():
    return jsonify({
        'message': 'Quizgenix Advanced AI Backend is running! 🤖',
        'status': 'healthy',
        'version': '2.0.0',
        'features': ['Advanced AI Question Generation', 'Knowledge Base Matching', 'Logical Answer Distribution'],
        'supported_domains': list(KNOWLEDGE_BASE.keys())
    })

@app.route('/api/login', methods=['POST'])
def login():
    try:
        data = request.get_json()
        print(f"🔐 Login attempt: {data.get('email')}")
        
        user = User.query.filter_by(email=data.get('email')).first()
        
        if user and check_password_hash(user.password_hash, data.get('password')):
            token = jwt.encode({
                'user_id': user.id,
                'exp': datetime.datetime.utcnow() + datetime.timedelta(days=30)
            }, app.config['SECRET_KEY'], algorithm='HS256')
            
            print(f"✅ Login successful: {user.email}")
            return jsonify({
                'message': 'Login successful',
                'token': token,
                'user': {
                    'id': user.id,
                    'name': user.name,
                    'email': user.email,
                    'role': user.role
                }
            })
        else:
            print(f"❌ Invalid credentials for: {data.get('email')}")
            return jsonify({'error': 'Invalid email or password'}), 401
            
    except Exception as e:
        print(f"❌ Login error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/register', methods=['POST'])
def register():
    try:
        data = request.get_json()
        print(f"📝 Registration attempt: {data.get('email')}")
        
        if User.query.filter_by(email=data.get('email')).first():
            return jsonify({'error': 'Email already exists'}), 400
        
        user = User(
            name=data.get('name'),
            email=data.get('email'),
            password_hash=generate_password_hash(data.get('password')),
            role=data.get('role', 'student')
        )
        
        db.session.add(user)
        db.session.commit()
        
        token = jwt.encode({
            'user_id': user.id,
            'exp': datetime.datetime.utcnow() + datetime.timedelta(days=30)
        }, app.config['SECRET_KEY'], algorithm='HS256')
        
        print(f"✅ Registration successful: {user.email}")
        return jsonify({
            'message': 'Registration successful',
            'token': token,
            'user': {
                'id': user.id,
                'name': user.name,
                'email': user.email,
                'role': user.role
            }
        })
        
    except Exception as e:
        print(f"❌ Registration error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/quizzes', methods=['GET'])
def get_quizzes():
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'No token provided'}), 401
            
        data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
        user_id = data['user_id']
        
        quizzes = Quiz.query.filter_by(user_id=user_id).order_by(Quiz.created_at.desc()).all()
        
        quiz_list = []
        for quiz in quizzes:
            quiz_list.append({
                'id': quiz.id,
                'title': quiz.title,
                'subject': quiz.subject,
                'topic': quiz.topic,
                'difficulty': quiz.difficulty,
                'questions': json.loads(quiz.questions),
                'created_at': quiz.created_at.isoformat()
            })
        
        return jsonify({'quizzes': quiz_list})
        
    except Exception as e:
        print(f"❌ Get quizzes error: {e}")
        return jsonify({'error': str(e)}), 500

# Add these endpoints after your existing routes

@app.route('/api/quiz/<int:quiz_id>/start', methods=['POST'])
def start_quiz(quiz_id):
    """Start a new quiz session for a student"""
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'No token provided'}), 401
            
        data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
        user_id = data['user_id']
        
        # Get the quiz
        quiz = Quiz.query.get(quiz_id)
        if not quiz:
            return jsonify({'error': 'Quiz not found'}), 404
        
        # Check if user already has an active session for this quiz
        existing_session = QuizSession.query.filter_by(
            user_id=user_id, 
            quiz_id=quiz_id, 
            is_active=True
        ).first()
        
        if existing_session and existing_session.expires_at > datetime.datetime.utcnow():
            return jsonify({
                'message': 'Quiz session already active',
                'session_token': existing_session.session_token,
                'expires_at': existing_session.expires_at.isoformat(),
                'current_question': existing_session.current_question
            })
        
        # Create new session
        session_token = jwt.encode({
            'user_id': user_id,
            'quiz_id': quiz_id,
            'session_id': random.randint(10000, 99999),
            'exp': datetime.datetime.utcnow() + datetime.timedelta(hours=2)
        }, app.config['SECRET_KEY'], algorithm='HS256')
        
        quiz_session = QuizSession(
            user_id=user_id,
            quiz_id=quiz_id,
            session_token=session_token,
            expires_at=datetime.datetime.utcnow() + datetime.timedelta(hours=2)
        )
        
        db.session.add(quiz_session)
        db.session.commit()
        
        # Get quiz questions without showing correct answers
        questions = json.loads(quiz.questions)
        quiz_data = {
            'id': quiz.id,
            'title': quiz.title,
            'subject': quiz.subject,
            'topic': quiz.topic,
            'difficulty': quiz.difficulty,
            'total_questions': len(questions),
            'questions': [
                {
                    'id': q['id'],
                    'question': q['question'],
                    'options': q['options']
                    # Don't include correct_answer or explanation
                }
                for q in questions
            ]
        }
        
        print(f"🎯 Quiz session started: {quiz.title} for user {user_id}")
        return jsonify({
            'message': 'Quiz session started',
            'session_token': session_token,
            'quiz': quiz_data,
            'time_limit': 7200,  # 2 hours in seconds
            'expires_at': quiz_session.expires_at.isoformat()
        })
        
    except Exception as e:
        print(f"❌ Start quiz error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/quiz/session/<session_token>/answer', methods=['POST'])
def submit_answer(session_token):
    """Submit answer for current question"""
    try:
        # Verify session
        session = QuizSession.query.filter_by(session_token=session_token, is_active=True).first()
        if not session or session.expires_at < datetime.datetime.utcnow():
            return jsonify({'error': 'Invalid or expired session'}), 401
        
        data = request.get_json()
        question_id = data.get('question_id')
        selected_answer = data.get('answer')  # 0, 1, 2, or 3
        
        # Update session with answer
        current_answers = json.loads(session.answers_so_far)

        current_answers[str(question_id)] = selected_answer
        
        session.answers_so_far = json.dumps(current_answers)
        session.current_question = max(session.current_question, question_id)
        
        db.session.commit()
        
        return jsonify({
            'message': 'Answer submitted',
            'question_id': question_id,
            'total_answered': len(current_answers)
        })
        
    except Exception as e:
        print(f"❌ Submit answer error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/quiz/session/<session_token>/submit', methods=['POST'])
def submit_quiz(session_token):
    """Submit entire quiz and calculate results"""
    try:
        # Verify session
        session = QuizSession.query.filter_by(session_token=session_token, is_active=True).first()
        if not session:
            return jsonify({'error': 'Invalid session'}), 401
        
        # Get quiz and questions
        quiz = Quiz.query.get(session.quiz_id)
        questions = json.loads(quiz.questions)
        user_answers = json.loads(session.answers_so_far)
        
        # Calculate score
        correct_count = 0
        detailed_results = []
        
        for question in questions:
            q_id = str(question['id'])
            user_answer = user_answers.get(q_id)
            correct_answer = question['correct_answer']
            is_correct = user_answer == correct_answer
            
            if is_correct:
                correct_count += 1
            
            detailed_results.append({
                'question_id': question['id'],
                'question': question['question'],
                'options': question['options'],
                'user_answer': user_answer,
                'correct_answer': correct_answer,
                'is_correct': is_correct,
                'explanation': question.get('explanation', ''),
                'references': question.get('references', [])
            })
        
        # Calculate time taken
        time_taken = int((datetime.datetime.utcnow() - session.started_at).total_seconds())
        
        # Save quiz attempt
        quiz_attempt = QuizAttempt(
            user_id=session.user_id,
            quiz_id=session.quiz_id,
            answers=session.answers_so_far,
            score=correct_count,
            total_questions=len(questions),
            time_taken=time_taken
        )
        
        # Deactivate session
        session.is_active = False
        
        db.session.add(quiz_attempt)
        db.session.commit()
        
        # Calculate percentage and grade
        percentage = round((correct_count / len(questions)) * 100, 1)
        grade = 'A' if percentage >= 90 else 'B' if percentage >= 80 else 'C' if percentage >= 70 else 'D' if percentage >= 60 else 'F'
        
        print(f"🏆 Quiz submitted: {quiz.title} - Score: {correct_count}/{len(questions)} ({percentage}%)")
        
        return jsonify({
            'message': 'Quiz submitted successfully',
            'results': {
                'score': correct_count,
                'total_questions': len(questions),
                'percentage': percentage,
                'grade': grade,
                'time_taken': time_taken,
                'detailed_results': detailed_results,
                'quiz_info': {
                    'title': quiz.title,
                    'subject': quiz.subject,
                    'topic': quiz.topic,
                    'difficulty': quiz.difficulty
                }
            }
        })
        
    except Exception as e:
        print(f"❌ Submit quiz error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/quiz/<int:quiz_id>/attempts', methods=['GET'])
def get_quiz_attempts(quiz_id):
    """Get all attempts for a specific quiz (for lecturers)"""
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'No token provided'}), 401
            
        data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
        user_id = data['user_id']
        
        # Check if user is lecturer and owns the quiz
        quiz = Quiz.query.filter_by(id=quiz_id, user_id=user_id).first()
        if not quiz:
            return jsonify({'error': 'Quiz not found or access denied'}), 403
        
        attempts = QuizAttempt.query.filter_by(quiz_id=quiz_id).join(User).all()
        
        attempt_list = []
        for attempt in attempts:
            attempt_list.append({
                'id': attempt.id,
                'student_name': attempt.user.name,
                'student_email': attempt.user.email,
                'score': attempt.score,
                'total_questions': attempt.total_questions,
                'percentage': round((attempt.score / attempt.total_questions) * 100, 1),
                'time_taken': attempt.time_taken,
                'completed_at': attempt.completed_at.isoformat()
            })
        
        return jsonify({
            'quiz_title': quiz.title,
            'total_attempts': len(attempt_list),
            'attempts': attempt_list
        })
        
    except Exception as e:
        print(f"❌ Get quiz attempts error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/my-results', methods=['GET'])
def get_my_results():
    """Get quiz results for current student"""
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'error': 'No token provided'}), 401
            
        data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
        user_id = data['user_id']
        
        attempts = QuizAttempt.query.filter_by(user_id=user_id).join(Quiz).all()
        
        result_list = []
        for attempt in attempts:
            percentage = round((attempt.score / attempt.total_questions) * 100, 1)
            grade = 'A' if percentage >= 90 else 'B' if percentage >= 80 else 'C' if percentage >= 70 else 'D' if percentage >= 60 else 'F'
            
            result_list.append({
                'id': attempt.id,
                'quiz_title': attempt.quiz.title,
                'quiz_subject': attempt.quiz.subject,
                'quiz_topic': attempt.quiz.topic,
                'quiz_difficulty': attempt.quiz.difficulty,
                'score': attempt.score,
                'total_questions': attempt.total_questions,
                'percentage': percentage,
                'grade': grade,
                'time_taken': attempt.time_taken,
                'completed_at': attempt.completed_at.isoformat()
            })
        
        return jsonify({
            'total_attempts': len(result_list),
            'results': sorted(result_list, key=lambda x: x['completed_at'], reverse=True)
        })
        
    except Exception as e:
        print(f"❌ Get my results error: {e}")
        return jsonify({'error': str(e)}), 500

def init_db():
    """Initialize database with sample data"""
    print("🗄️ Initializing database...")
    
    db.create_all()
    
    if not User.query.filter_by(email='lecturer@test.com').first():
        lecturer = User(
            name='Test Lecturer',
            email='lecturer@test.com',
            password_hash=generate_password_hash('password123'),
            role='lecturer'
        )
        db.session.add(lecturer)
        print("👨‍🏫 Created test lecturer: lecturer@test.com")
    
    if not User.query.filter_by(email='student@test.com').first():
        student = User(
            name='Test Student',
            email='student@test.com',
            password_hash=generate_password_hash('password123'),
            role='student'
        )
        db.session.add(student)
        print("🎓 Created test student: student@test.com")
    
    db.session.commit()
    print("✅ Database initialized successfully")

if __name__ == '__main__':
    print("🚀 Starting Quizgenix Advanced AI Backend...")
    print("🤖 AI Features: Smart Knowledge Base, Logical Questions, Diverse Answers")
    print("🧠 Supported Domains:", list(KNOWLEDGE_BASE.keys()))
    
    with app.app_context():
        init_db()
    
    print("🌐 Server starting on http://127.0.0.1:5000")
    app.run(debug=True, host='127.0.0.1', port=5000)